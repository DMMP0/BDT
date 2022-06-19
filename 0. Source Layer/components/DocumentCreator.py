from unicodedata import name
import pandas as pd
import random
from random import randint
import uuid
import numpy as np
import datetime
import docx
import re

from docx.shared import Pt, Mm

temp = []
fiscale = []
users = []

threshold = 50  # this is the number of companies request for the bank data

business_data = pd.read_csv('../../assests/0. Source Data/credit data/company_information.csv')
personal_data = pd.read_csv('../../assests/0. Source Data/credit data/user_information.csv')

business_data = pd.DataFrame(data=business_data, columns=business_data.columns)
business_data = business_data.iloc[:, 1:]


def df_to_word(data, report_name: str) -> docx.Document:
    # assert type(data) == dict, 'data has to be dict'
    doc = docx.Document()
    section = doc.sections[0]
    section.left_margin = Mm(0)
    section.right_margin = Mm(0)
    section.top_margin = Mm(10)
    section.bottom_margin = Mm(10)

    data = pd.DataFrame(data)  # My input data is in the 2D list form
    table = doc.add_table(rows=(data.shape[0]), cols=data.shape[1])  # First row are table headers!
    table.allow_autofit = True
    table.autofit = True

    table = doc.add_table(data.shape[0] + 1, data.shape[1])
    for i, column in enumerate(data):
        for row in range(data.shape[0]):
            table.cell(row, i).text = str(data[column][row])
    doc.save(f'./{report_name}')


def to_pdf(data):
    data = pd.DataFrame(data)
    data.to_html('./reports/report.html')


def to_rtf(data):
    data = pd.DataFrame(data)
    ## read in the text as a string
    with open('./reports/report.txt', 'a') as f:
        dfAsString = data.to_string(header=False, index=False)
        f.write(dfAsString)


def to_tex(data):
    data = pd.DataFrame(data)
    data.to_latex('./reports/report.tex')


def to_excel(data):
    data = pd.DataFrame(data)
    data.to_excel('./reports/report.xlsx')


def to_clipboard(data):
    data = pd.DataFrame(data)
    data.to_clipboard('./reports/report.clp')


for i in range(0, 6):
    data = business_data.iloc[0:threshold]
    rnd = randint(10, threshold)
    idx = rnd
    print("Randome number", rnd)
    df = data.take(np.random.permutation(len(data))[:idx])
    if i == 0:
        to_rtf(df)
        print(idx)
    if i == 1:
        df_to_word(data[0:idx], './reports/report.doc')
        print(idx)
    if i == 2:
        to_pdf(df)
        print(idx)
    if i == 3:
        to_tex(df)
        print(idx)
    if i == 4:
        to_excel(df)
        print(idx)
    if i == 5:
        to_clipboard(df)
        print(idx)

# df_to_word(data, 'report.doc')
# df_to_word(data, 'report.docx')
# to_pdf(data)
# to_rtf(data)
# to_tex(data)
# to_excel(data)
# to_clipboard(data)

# print(data)

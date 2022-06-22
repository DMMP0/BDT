
import pandas as pd

import io
import csv
from docx import Document



bank_columns = ['registration_number', 'name', 'established_date', 'country', 'number_of_employees', 'purpose',
           'phone_number', 'email',
           'bank_name', 'bank_country', 'open_new_credit_in_6_months', 'amount_in_6_months', 'new_credit_in_12_months',
           'new_credit_in_18_months',
           'amount_in_12_months', 'amount_in_18_months', 'house_mortgage', 'amount_of_house_mortgage',
           'amount_due_mortgage', 'house_property',
           'total_house_amount', 'credit_card_number', 'actual_debit_credit_cards', 'monthly_income', 'savings',
           'other_savings']


questura_columns = ['registration_number','name','established_date','country','number_of_employees','purpose',
'phone_number','email','questura_country','bankruptcy','inscred','fraudis','investigation','accused','condamned','civ_pass']


broker_columns = ['registration_number', 'name', 'established_date', 'country', 'number_of_employees', 'purpose',
           'phone_number', 'email','agencey_country','agency_name','debit_id','installment','installment_ammount','agency_country']

def read_docx_tables(filename, tab_id=None, **kwargs):
    """
    parse table(s) from a Word Document (.docx) into Pandas DataFrame(s)

    Parameters:
        filename:   file name of a Word Document

        tab_id:     parse a single table with the index: [tab_id] (counting from 0).
                    When [None] - return a list of DataFrames (parse all tables)

        kwargs:     arguments to pass to `pd.read_csv()` function

    Return: a single DataFrame if tab_id != None or a list of DataFrames otherwise
    """

    def read_docx_tab(tab, **kwargs):
        vf = io.StringIO()
        writer = csv.writer(vf)
        for row in tab.rows:
            writer.writerow(cell.text for cell in row.cells)
        vf.seek(0)
        return pd.read_csv(vf, **kwargs)

    doc = Document(filename)
    if tab_id is None:
        for tab in doc.tables:
            df = read_docx_tab(tab, **kwargs)
        return df
    else:
        try:
            return read_docx_tab(doc.tables[tab_id], **kwargs)
        except IndexError:
            print('Error: specified [tab_id]: {}  does not exist.'.format(tab_id))
            raise


# print(len(excel))
# read excel
def excel_to_dict(filepath: str) -> dict:
    """The function returns a dictionary from a filepath. The dictionary will have the filename as key and the values
    as another dictionary"""
    # global columns
    # Read a text file to a dataframe using read_table function

    data = pd.read_excel(filepath)
    data = pd.DataFrame(data)
    # columns = data.columns
    name = str(filepath).split('/')[-1]
    return {name: data.to_dict(orient='index')}
    # temp = data.to_json('./jsons/' + name + '.json', indent=4, orient='index')


def html_to_dict(filepath: str) -> dict:
    """The function returns a dictionary from a filepath. The dictionary will have the filename as key and the values
        as another dictionary"""
    # Read a text file to a dataframe using read_table function
    name = filepath.split('/')[-1]
    data = pd.read_html(filepath)
    return {name: data[0].to_dict(orient='index')}
    # data[0].to_json('./jsons/' + name + '.json', indent=4, orient='index')


def docx_to_dict(filepath: str) -> dict:
    """The function returns a dictionary from a filepath. The dictionary will have the filename as key and the values
        as another dictionary"""
    global bank_columns
    global questura_columns
    global broker_columns
    name = filepath.split('/')[-1]
    df = read_docx_tables(filepath)
    df = pd.DataFrame(df)

    if 'Questura' in name:
        columns = questura_columns
    elif 'Broker' in name:
        columns = broker_columns
    else:
        columns = bank_columns

    df = pd.DataFrame(df.values,
                      columns=columns)  # NB: don't modify
    return {name: df.to_dict(orient='index')}
    # df.to_json('./jsons/' + name + '.json', indent=4, orient='index')


def txt_to_dict(filepath: str) -> dict:
    """The function returns a dictionary from a filepath. The dictionary will have the filename as key and the values
           as another dictionary"""
    # Read a text file to a dataframe using read_table function
    name = str(filepath).split('/')[-1]
    data = pd.read_csv(filepath, sep='\t')
    data = pd.DataFrame(data)
    return {name: data.to_dict(orient='index')}
    # temp = data.to_json('./jsons/' + name + '.json', indent=4, orient='index')

import os
import random
from random import randint

import ccard
import docx
import pandas as pd

val = 3000.00
source = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))  # finds the 0. Data Source directory

class Word:

    def __init__(self):
        self.data = []
        self.bank = []

    def clean(self):
        self.data = []
        self.bank = []

    def set_data(self, data):
        self.data = data

    def set_bank(self, bank):
        self.bank = bank

    def create_bank_data(self, bdata):
        bank_data = bdata
        banksName = str(self.bank[0])
        banksName = banksName.replace(' ', '_')
        banksName = banksName.replace('/', '_')
        bank_data = pd.DataFrame(bank_data)
        bank_data['bank_name'] = self.bank[0]
        bank_data['bank_country'] = self.bank[1]
        bank_data['open_new_credit_in_6_months'] = 0
        bank_data['ammount_in_6_months'] = 0.00
        bank_data['new_credit_in_12_months'] = 0.00
        bank_data['new_credit_in_18_months'] = 0.00
        bank_data['ammount_in_12_months'] = 0.00
        bank_data['ammount_in_18_months'] = 0.00
        bank_data['house_mortage'] = 1
        bank_data['amount_of_house_mortage'] = 0.00
        bank_data['amount_duee_mortage'] = 0.00
        bank_data['house_property'] = 00
        bank_data['total_house_amount'] = 0.00
        bank_data['credit_card_number'] = 0
        bank_data['actual_debit_credit_cards'] = 0.00
        bank_data['monthly_income'] = 0.00
        bank_data['savings'] = 0.00
        bank_data['other_savings'] = 0.00

        return bank_data

    def create_questura_data(self, questuradata) -> any:
        questura_data = questuradata
        # print(bank['names'])
        # print(self.bank[0])

        # bank = pd.DataFrame(bank)
        # bank = pd.DataFrame(data = bank.values , columns=bank.columns)
        # bank = list(bank)
        # print(bank)
        banksName = str(self.bank[0])
        banksName = banksName.replace(' ', '_')
        banksName = banksName.replace('/', '_')
        questura_data = pd.DataFrame(questura_data)
        questura_data['questura_country'] = self.bank[1]
        questura_data['bankruptcy'] = bool
        questura_data['inscred'] = "${:,.2f}".format(val + random.randint(5000, 100_000_000_000))
        questura_data['fraudis'] = bool
        questura_data['investegation'] = bool
        questura_data['accused'] = bool
        questura_data['condamned'] = bool
        questura_data['civ_pass'] = bool

        return questura_data

    def create_risk_broker_data(self, broker_data) -> any:
        broker_data = broker_data
        banksName = str(self.bank[0])
        banksName = banksName.replace(' ', '_')
        banksName = banksName.replace('/', '_')
        broker_data = pd.DataFrame(broker_data)
        broker_data['agencey_country'] = self.bank[1]
        broker_data['agency_name'] = ''
        broker_data['from30to60'] = random.randint(0, 3)
        broker_data['from60to90'] = random.randint(0, 3)
        broker_data['morethan90'] = random.randint(0, 3)
        broker_data['debit_id'] = ''
        broker_data['insolvent'] = bool
        broker_data['insolvent_ammount'] = "${:,.2f}".format(val + random.randint(5000, 100000000000))

    def bank_to_word(self):
        bank_data = pd.DataFrame(self.data)
        banksName = self.bank['names']
        bank_data['bank_name'] = self.bank['names']
        bank_data['bank_country'] = self.bank['country']
        banksName = banksName.replace(' ', '_')
        banksName = banksName.replace('/', '_')
        bank_data = self.create_bank_data(bank_data)
        banksName = str(self.bank[0])
        banksName = banksName.replace(' ', '_')
        banksName = banksName.replace('/', '_')
        bank_data = pd.DataFrame(bank_data)
        for index, item in bank_data.iterrows():
            bank_data.loc[index, 'open_new_credit_in_6_months'] = random.randint(0, 10)
            bank_data.loc[index, 'ammount_in_6_months'] = "${:,.2f}".format(val + random.randint(5000, 100000000000))
            bank_data.loc[index, 'new_credit_in_12_months'] = random.randint(0, 3)
            bank_data.loc[index, 'new_credit_in_18_months'] = random.randint(0, 3)
            bank_data.loc[index, 'ammount_in_12_months'] = "${:,.2f}".format(val + random.randint(2000, 20000))
            bank_data.loc[index, 'ammount_in_18_months'] = "${:,.2f}".format(val + random.randint(5000, 100000000000))
            bank_data.loc[index, 'house_mortage'] = random.choice([True, False])
            if (bank_data.loc[index, 'house_mortage'] == True):
                bank_data.loc[index, 'amount_of_house_mortage'] = "${:,.2f}".format(val + random.randint(50000, 500000))
                bank_data.loc[index, 'amount_duee_mortage'] = "${:,.2f}".format(val + random.randint(50000, 500000))
                bank_data.loc[index, 'house_property'] = random.choice([True, False])
            else:
                bank_data.loc[index, 'house_property'] = "${:,.2f}".format(val + random.randint(0, 500000))
            if (bank_data.loc[index, 'house_property'] is True):
                bank_data.loc[index, 'total_house_amount'] = "${:,.2f}".format(val + random.randint(500, 50000))
            bank_data.loc[index, 'credit_card_number'] = random.randint(1, 5)
            bank_data.loc[index, 'credit_card_limit_total'] = random.randint(10000, 500000)
            bank_data.loc[index, 'actual_debit_credit_cards'] = random.randint(0, 5)
            bank_data.loc[index, 'monthly_income'] = "${:,.2f}".format(val + random.randint(0, 50000))
            bank_data.loc[index, 'savings'] = "${:,.2f}".format(val + random.randint(1000, 100000))
            bank_data.loc[index, 'other_savings'] = "${:,.2f}".format(val + random.randint(100, 20000))
        doc = docx.Document()
        table = doc.add_table(rows=bank_data.shape[0], cols=bank_data.shape[1])
        table_cells = table._cells
        for i in range(bank_data.shape[0]):
            for j in range(bank_data.shape[1]):
                table_cells[j + i * bank_data.shape[1]].text = str(bank_data.values[i][j])
                doc.save(source+'/components/reports/'+banksName+'(Bank).docx')

    def police_to_word(self):
        questura_data = pd.DataFrame(self.data)
        banksName = self.bank['names']
        # questura_data['bank_name'] = self.bank['names']
        # questura_data['bank_country'] = self.bank['country']
        # print(banksName)
        banksName = banksName.replace(' ', '_')
        banksName = banksName.replace('/', '_')
        bank = self.bank
        questura_data = self.create_questura_data(questura_data)
        banksName = str(self.bank[0])
        banksName = banksName.replace(' ', '_')
        banksName = banksName.replace('/', '_')
        questura_data = pd.DataFrame(questura_data)
        # print(banksName)
        for index, item in questura_data.iterrows():
            questura_data = pd.DataFrame(questura_data)
            questura_data.loc[index, 'questura_country'] = self.bank[1]
            questura_data.loc[index, 'bankruptcy'] = random.choice([True, False])
            if questura_data.loc[index, 'bankruptcy']:
                questura_data.loc[index, 'inscred'] = "${:,.2f}".format(val + random.randint(500000, 1000000))
            if not questura_data.loc[index, 'bankruptcy']:
                questura_data.loc[index, 'fraudis'] = False
            else:
                questura_data.loc[index, 'fraudis'] = True
            questura_data.loc[index, 'investegation'] = random.choice([True, False])
            questura_data.loc[index, 'accused'] = random.choice([True, False])
            questura_data.loc[index, 'condamned'] = random.choice([True, False])
            questura_data.loc[index, 'civ_pass'] = random.choice([True, False])
        doc = docx.Document()
        table = doc.add_table(rows=questura_data.shape[0], cols=questura_data.shape[1])
        table_cells = table._cells
        for i in range(questura_data.shape[0]):
            for j in range(questura_data.shape[1]):
                table_cells[j + i * questura_data.shape[1]].text = str(questura_data.values[i][j])
        doc.save(source+'/components/reports/'+banksName+'(Questura).docx')

    agencey_names = ['CRIF', 'CTC', 'Banca d italia', 'experian']

    def broker_to_word(self):
        broker_data = pd.DataFrame(self.data)
        banksName = self.bank['names']
        banksName = banksName.replace(' ', '_')
        banksName = banksName.replace('/', '_')
        broker_data = self.create_risk_broker_data(broker_data)
        banksName = self.bank[0]
        banksName = banksName.replace(' ', '_')
        banksName = banksName.replace('/', '_')
        broker_data = pd.DataFrame(broker_data)
        # print(banksName)
        for index, item in broker_data.iterrows():
            broker_data = pd.DataFrame(broker_data)
            broker_data.loc[index, 'agency_country'] = self.bank[1]
            broker_data.loc[index, 'agency_name'] = self.agencey_names[(randint(0, 3))]
            broker_data.loc[index, 'debit_id'] = str(ccard.americanexpress())
            broker_data.loc[index, 'installment'] = random.choice([True, False])
            broker_data.loc[index, 'installment_amount'] = "${:,.2f}".format(val + random.randint(5000, 100000000000))

        doc = docx.Document()
        table = doc.add_table(rows=broker_data.shape[0], cols=broker_data.shape[1])
        table_cells = table._cells
        for i in range(broker_data.shape[0]):
            for j in range(broker_data.shape[1]):
                table_cells[j + i * broker_data.shape[1]].text = str(broker_data.values[i][j])
        doc.save(source+'/components/reports/'+banksName+'(Broker).docx')
